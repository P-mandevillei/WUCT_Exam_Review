# Vibed by Gemini 3.5
import re
import io
import copy
import docx
from docx import Document
from docx.text.paragraph import Paragraph
from docx.table import Table
from docx.shared import RGBColor

# Regexes for parsing
GRADESCOPE_Q_REGEX = re.compile(
    r'Question\s+(\d+(?:[a-zA-Z]|\.[a-zA-Z0-9]+)?(?:\([ivxIVX]+\))?)(?:\s+|$|\()',
    re.IGNORECASE
)

def int_to_roman(num):
    """
    Helper to convert integer to lowercase Roman numeral
    """
    val = [10, 9, 5, 4, 1]
    syb = ["x", "ix", "v", "iv", "i"]
    roman_num = ''
    i = 0
    while num > 0:
        for _ in range(num // val[i]):
            roman_num += syb[i]
            num -= val[i]
        i += 1
    return roman_num

def extract_valid_q_ids(question_names):
    """
    Extracts normalized question IDs (e.g., '1', '2a', '2a(i)') from Gradescope question names.
    """
    valid_ids = set()
    for name in question_names:
        m = GRADESCOPE_Q_REGEX.search(name)
        if m:
            valid_ids.add(m.group(1).lower())
    return valid_ids

def iter_elements(document):
    """
    Generate each paragraph and table child within the document body,
    in document order.
    """
    for child in document.element.body.iterchildren():
        tag = child.tag
        if tag.endswith('}p'):
            yield Paragraph(child, document)
        elif tag.endswith('}tbl'):
            yield Table(child, document)

def delete_paragraph(paragraph):
    """
    Safely deletes a paragraph from the document XML tree.
    """
    p = paragraph._p
    p.getparent().remove(p)
    p._p = p._element = None

def parse_docx(file_stream, valid_q_ids):
    """
    Parses the uploaded DOCX document and groups elements by question ID.
    Decodes the implicit hierarchy based on visual left indentation.
    """
    doc = Document(file_stream)
    
    questions = {}
    current_problem = None
    level_0_counter = 0
    level_1_counter = 0
    current_q_id = None
    
    # Match problem headings at the start of a paragraph (case-insensitive, style-agnostic)
    prob_regex = re.compile(r'^\s*Problem\s*#?\s*(\d+)', re.IGNORECASE)
    
    # Patterns indicating grading rubric/answer key rather than a question prompt
    rubric_patterns = [
        r'^\s*[\+\-\(]*\d+\s*points?\b',
        r'^\s*\+\s*\d+',
        r'^\s*Fully\s+correct',
        r'^\s*Partially\s+correct',
        r'^\s*\(?\s*max\s+\d+\s+points?',
        r'^\s*Note\s*\:',
    ]
    
    intro_elements = []
    
    for element in iter_elements(doc):
        if isinstance(element, Paragraph):
            text = element.text.strip()
            
            # Check problem title heading
            prob_match = prob_regex.match(text)
            if prob_match:
                current_problem = int(prob_match.group(1))
                level_0_counter = 0
                level_1_counter = 0
                current_q_id = None
                intro_elements.append(element)
                continue
                
            if current_problem is not None:
                pPr = element._p.get_or_add_pPr()
                numPr = pPr.numPr
                
                # Identify if the text is grading rubric or point values
                is_rubric = any(re.search(pat, text, re.IGNORECASE) for pat in rubric_patterns)
                
                # Ignore list markers for empty paragraphs and rubric paragraphs
                if numPr is not None and text != "" and not is_rubric:
                    # Get Left Indent
                    ilvl = numPr.ilvl.val if numPr.ilvl is not None else 0
                    left_indent = element.paragraph_format.left_indent
                    if left_indent:
                        L = left_indent.inches
                    else:
                        # Fallback using ilvl
                        L = 0.5 if ilvl == 0 else (1.0 if ilvl == 1 else 1.5)
                        
                    # Identify Level
                    if 0.3 <= L <= 0.7:
                        level_0_counter += 1
                        level_1_counter = 0
                        level_0_char = chr(ord('a') + level_0_counter - 1)
                        extracted_id = f"{current_problem}{level_0_char}"
                        if extracted_id in valid_q_ids:
                            current_q_id = extracted_id
                            questions[current_q_id] = []
                        else:
                            current_q_id = None
                    elif 0.8 <= L <= 1.2:
                        level_1_counter += 1
                        level_0_char = chr(ord('a') + level_0_counter - 1)
                        roman_val = int_to_roman(level_1_counter)
                        extracted_id = f"{current_problem}{level_0_char}({roman_val})"
                        if extracted_id in valid_q_ids:
                            current_q_id = extracted_id
                            questions[current_q_id] = []
                        else:
                            current_q_id = None
                    elif L > 1.2:
                        # Options, keep appending to active current_q_id
                        pass
        
        # Append element (Paragraph or Table) to appropriate list
        if current_q_id is not None:
            if current_q_id not in questions:
                questions[current_q_id] = []
            questions[current_q_id].append(element)
        else:
            intro_elements.append(element)
            
    return questions, intro_elements

def generate_segmented_report(file_stream, questions_dict, type_df):
    """
    Generates a new DOCX report by grouping parsed questions into sections by type.
    We load the original document first to inherit styles, relationships (images),
    headers, and footers, clear the body, and then build the new document structure.
    """
    # Load the original document to inherit styling and media relationships
    out_doc = Document(file_stream)
    
    # Clear all paragraphs and tables from the body of out_doc
    for p in list(out_doc.paragraphs):
        delete_paragraph(p)
    for t in list(out_doc.tables):
        tbl = t._tbl
        tbl.getparent().remove(tbl)
        
    # Main Report Title
    title_p = out_doc.add_paragraph()
    title_run = title_p.add_run("WUCT Exam Analysis - Question Type Report")
    title_p.style = 'Title'
    
    desc_p = out_doc.add_paragraph()
    desc_p.add_run(
        "This report reorganizes the exam questions into sections according to their performance metrics "
        "and classification types derived from GradeScope student scores."
    )
    desc_p.style = 'Normal'
    
    # Define the 5 types
    types_list = [
        ('type1', 'Type 1 - Easy for Both Divisions'),
        ('type2', 'Type 2 - Lower Division Discrimination'),
        ('type3', 'Type 3 - Higher Division Discrimination'),
        ('type4', 'Type 4 - Too Hard / Needs Review'),
        ('type5', 'Type 5 - Confusing (Easy for LD, Hard for HD)')
    ]
    
    # Map GradeScope question names to normalized IDs
    # e.g., "8: Question 2a(i) (1.0 pts)" -> "2a(i)"
    gradescope_to_norm = {}
    for g_q in type_df['question'].unique():
        m = GRADESCOPE_Q_REGEX.search(g_q)
        if m:
            gradescope_to_norm[g_q] = m.group(1).lower()
            
    # Re-group questions by type
    grouped_questions = {t[0]: [] for t in types_list}
    
    # Process type_df (include even unmatched questions)
    for _, row in type_df.iterrows():
        g_q = row['question']
        q_type = row['type']
        norm_id = gradescope_to_norm.get(g_q)
        if norm_id:
            if q_type in grouped_questions:
                if (norm_id, g_q) not in grouped_questions[q_type]:
                    grouped_questions[q_type].append((norm_id, g_q))
                
    # Any question in the parsed document that is not in any category goes to uncategorized if it exists
    categorized_norm_ids = set()
    for q_type, q_list in grouped_questions.items():
        if q_type != 'uncategorized':
            for norm_id, _ in q_list:
                categorized_norm_ids.add(norm_id)
                
    if 'uncategorized' in grouped_questions:
        for norm_id in questions_dict:
            if norm_id not in categorized_norm_ids:
                # Find its GradeScope name if available, otherwise use a placeholder name
                g_name = None
                for g, n in gradescope_to_norm.items():
                    if n == norm_id:
                        g_name = g
                        break
                if g_name is None:
                    g_name = f"Question {norm_id.upper()}"
                grouped_questions['uncategorized'].append((norm_id, g_name))
            
    # Write each type section to the document
    for q_type, title in types_list:
        questions_of_type = grouped_questions[q_type]
        
        # Add heading for the question type
        h = out_doc.add_paragraph()
        hrun = h.add_run(title)
        h.style = 'Heading 1'
        
        if not questions_of_type:
            empty_p = out_doc.add_paragraph()
            empty_run = empty_p.add_run("No questions fall into this category.")
            empty_run.italic = True
            continue
            
        for norm_id, g_q in questions_of_type:
            # Add a sub-heading for the question
            sh = out_doc.add_paragraph()
            shrun = sh.add_run(g_q)
            sh.style = 'Heading 2'
            
            # Append elements if matched, else write warning placeholder
            if norm_id in questions_dict and questions_dict[norm_id]:
                ref_el = sh._p
                elements = questions_dict[norm_id]
                for element in elements:
                    if isinstance(element, Paragraph):
                        p_elem = copy.deepcopy(element._p)
                        ref_el.addnext(p_elem)
                        ref_el = p_elem
                    elif isinstance(element, Table):
                        t_elem = copy.deepcopy(element._tbl)
                        ref_el.addnext(t_elem)
                        ref_el = t_elem
            else:
                # Add failed matching warning placeholder
                p_fail = out_doc.add_paragraph()
                run_fail = p_fail.add_run("Failed to match question content in the uploaded document.")
                run_fail.italic = True
                run_fail.font.color.rgb = RGBColor(180, 50, 50)  # Dark red / warning color
                
                # Insert immediately after subheading
                sh._p.addnext(p_fail._p)
                    
    # Save the document to an in-memory byte stream
    out_stream = io.BytesIO()
    out_doc.save(out_stream)
    out_stream.seek(0)
    return out_stream
